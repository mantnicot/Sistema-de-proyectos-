import io
import logging
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.config import settings
from app.infrastructure.export.export_utils import (
    build_content_sections,
    collect_excel_columns,
    format_field_value,
    get_oati_logo,
    get_universidad_logo,
    load_logo_raster,
)

logger = logging.getLogger(__name__)

BRAND_COLOR = colors.HexColor("#9B2D2D")
BRAND_RGB = RGBColor(155, 45, 45)
BRAND_HEX = "9B2D2D"


class DocumentExporter:
    """Generador de documentos con identidad institucional OATI."""

    def _header_meta_rows(self, project_name: str) -> list[list[str]]:
        return [
            ["", project_name.upper(), "Código:", ""],
            ["", f"Macroproceso: {settings.MACROPROCESO}", f"Versión: {settings.DOCUMENTO_VERSION}", ""],
            ["", f"Proceso: {settings.PROCESO}", "Fecha de Aprobación:", ""],
        ]

    def _add_word_logo_cell(self, cell, logo_path, fallback: str, width_inches: float = 0.9) -> None:
        cell.text = ""
        raster = load_logo_raster(logo_path)
        if raster:
            try:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(raster, width=Inches(width_inches))
                return
            except Exception as exc:
                logger.warning("No se pudo insertar logo Word: %s", exc)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(fallback)
        run.bold = True
        run.font.color.rgb = BRAND_RGB
        run.font.size = Pt(9)

    def _set_word_cell_shading(self, cell, fill_hex: str) -> None:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    def _add_word_header_table(self, doc: Document, project_name: str) -> None:
        logo_ud = get_universidad_logo()
        logo_oati = get_oati_logo()
        header = doc.add_table(rows=3, cols=4)
        header.style = "Table Grid"

        self._add_word_logo_cell(header.rows[0].cells[0], logo_ud, "UD", 0.85)
        header.rows[0].cells[1].text = project_name.upper()
        header.rows[0].cells[2].text = "Código:"
        self._add_word_logo_cell(header.rows[0].cells[3], logo_oati, "OATI", 1.1)
        header.rows[1].cells[1].text = f"Macroproceso: {settings.MACROPROCESO}"
        header.rows[1].cells[2].text = f"Versión: {settings.DOCUMENTO_VERSION}"
        header.rows[2].cells[1].text = f"Proceso: {settings.PROCESO}"
        header.rows[2].cells[2].text = "Fecha de Aprobación:"

        for row in header.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
                if cell.text in (project_name.upper(), "Código:", "Fecha de Aprobación:"):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = BRAND_RGB

    def _add_word_cover(self, doc: Document, project_name: str) -> None:
        logo_ud = get_universidad_logo()
        raster = load_logo_raster(logo_ud)
        if raster:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.add_run().add_picture(raster, width=Inches(1.5))
            except Exception:
                pass

        for text, size, bold in [
            (settings.INSTITUCION_NOMBRE, 16, True),
            ("FRANCISCO JOSÉ DE CALDAS", 12, True),
            (project_name.upper(), 14, True),
            (settings.OATI_NOMBRE, 12, True),
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Calibri"
            run.font.color.rgb = BRAND_RGB if bold else RGBColor(0, 0, 0)

    def _add_word_toc(self, doc: Document, sections: list[tuple[str, str]]) -> None:
        h = doc.add_paragraph()
        run = h.add_run("TABLA DE CONTENIDO")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = BRAND_RGB
        run.font.name = "Calibri"

        toc = doc.add_table(rows=len(sections) + 1, cols=2)
        toc.style = "Table Grid"
        toc.rows[0].cells[0].text = "Sección"
        toc.rows[0].cells[1].text = "Pág."
        self._set_word_cell_shading(toc.rows[0].cells[0], BRAND_HEX)
        self._set_word_cell_shading(toc.rows[0].cells[1], BRAND_HEX)
        for cell in toc.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True

        for i, (title, _) in enumerate(sections, start=1):
            toc.rows[i].cells[0].text = title
            toc.rows[i].cells[1].text = str(i + 2)
            toc.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_word_sections(self, doc: Document, sections: list[tuple[str, str]]) -> None:
        for title, content in sections:
            h = doc.add_paragraph()
            run = h.add_run(title)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = BRAND_RGB
            run.font.name = "Calibri"
            body = doc.add_paragraph(content)
            if body.runs:
                body.runs[0].italic = True
                body.runs[0].font.size = Pt(10)
                body.runs[0].font.name = "Calibri"

    def _pdf_logo_flowable(self, logo_path, fallback: str, w: float, h: float) -> Any:
        raster = load_logo_raster(logo_path)
        if raster:
            try:
                return Image(raster, width=w, height=h)
            except Exception:
                pass
        return Paragraph(
            f"<b>{fallback}</b>",
            ParagraphStyle("LogoFallback", fontSize=9, textColor=BRAND_COLOR, alignment=TA_CENTER),
        )

    def export_pdf(self, project: dict, fields: list[dict]) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
        styles = getSampleStyleSheet()
        story: list[Any] = []

        title_style = ParagraphStyle(
            "Title", parent=styles["Heading1"], alignment=TA_CENTER, fontSize=14, textColor=BRAND_COLOR, spaceAfter=12
        )
        subtitle_style = ParagraphStyle(
            "Subtitle", parent=styles["Normal"], alignment=TA_CENTER, fontSize=11, textColor=BRAND_COLOR, spaceAfter=10
        )
        section_style = ParagraphStyle(
            "Section", parent=styles["Heading2"], fontSize=11, textColor=BRAND_COLOR, spaceBefore=12, spaceAfter=6
        )
        body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, alignment=TA_LEFT)

        logo_ud = get_universidad_logo()
        logo_oati = get_oati_logo()
        meta = self._header_meta_rows(project["name"])

        header_cells = [
            [
                self._pdf_logo_flowable(logo_ud, "UD", 1.8 * cm, 1.8 * cm),
                Paragraph(f"<b>{meta[0][1]}</b>", body_style),
                Paragraph(f"<b>{meta[0][2]}</b>", body_style),
                self._pdf_logo_flowable(logo_oati, "OATI", 2.2 * cm, 1.2 * cm),
            ],
            ["", Paragraph(meta[1][1], body_style), Paragraph(meta[1][2], body_style), ""],
            ["", Paragraph(meta[2][1], body_style), Paragraph(meta[2][2], body_style), ""],
        ]
        header_table = Table(header_cells, colWidths=[3 * cm, 9 * cm, 4 * cm, 3 * cm])
        header_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F5E6E6")),
                    ("SPAN", (0, 1), (0, 2)),
                    ("SPAN", (3, 1), (3, 2)),
                ]
            )
        )
        story.append(header_table)
        story.append(Spacer(1, 0.8 * cm))

        cover_logo = load_logo_raster(logo_ud)
        if cover_logo:
            try:
                story.append(Image(cover_logo, width=3.5 * cm, height=3.5 * cm, hAlign="CENTER"))
                story.append(Spacer(1, 0.4 * cm))
            except Exception:
                pass

        story.append(Paragraph(settings.INSTITUCION_NOMBRE, title_style))
        story.append(Paragraph("FRANCISCO JOSÉ DE CALDAS", subtitle_style))
        story.append(Spacer(1, 0.8 * cm))
        story.append(Paragraph(project["name"].upper(), title_style))
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(settings.OATI_NOMBRE, subtitle_style))
        story.append(PageBreak())

        sections = build_content_sections(project, fields)

        story.append(Paragraph("TABLA DE CONTENIDO", section_style))
        toc_data = [["Sección", "Pág."]]
        for i, (title, _) in enumerate(sections, start=1):
            toc_data.append([title, str(i + 2)])
        toc_table = Table(toc_data, colWidths=[14 * cm, 2 * cm])
        toc_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), BRAND_COLOR),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(toc_table)
        story.append(PageBreak())

        for title, content in sections:
            story.append(Paragraph(f"<b>{title}</b>", section_style))
            story.append(Paragraph(f"<i>{content}</i>", body_style))
            story.append(Spacer(1, 0.25 * cm))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    def export_word(self, project: dict, fields: list[dict]) -> bytes:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        self._add_word_header_table(doc, project["name"])
        doc.add_paragraph()
        self._add_word_cover(doc, project["name"])
        doc.add_page_break()

        sections = build_content_sections(project, fields)
        self._add_word_toc(doc, sections)
        doc.add_page_break()
        self._add_word_sections(doc, sections)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_excel_all(self, projects: list[dict], fields: list[dict]) -> bytes:
        wb = Workbook()
        ws = wb.active
        ws.title = "Proyectos OATI"

        header_fill = PatternFill(start_color=BRAND_HEX, end_color=BRAND_HEX, fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        columns = collect_excel_columns(projects, fields)
        headers = ["ID", "Nombre", "Descripción"] + [label for _, label in columns]
        ws.append(headers)
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", wrap_text=True)

        for proj in projects:
            data = proj.get("data") or {}
            row = [
                proj.get("id", "N/A"),
                proj.get("name") or "N/A",
                format_field_value(proj.get("description"), "N/A"),
            ]
            for key, _ in columns:
                row.append(format_field_value(data.get(key, ""), "N/A"))
            ws.append(row)

        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max(max_len + 4, 12), 50)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.read()
